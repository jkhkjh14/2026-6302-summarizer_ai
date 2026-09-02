import os
from dotenv import load_dotenv

# Загружаем переменные из .env
load_dotenv()

from flask import Flask, request, jsonify, render_template, send_file
import tempfile
import PyPDF2
import docx
from newspaper import Article
import easyocr
from PIL import Image
import logging
import requests
from bs4 import BeautifulSoup
import json
import io
from docx import Document as DocxDocument
import re

app = Flask(__name__)
logging.basicConfig(level=logging.INFO)

# ==========================================
# Загружаем данные из переменных окружения
# ==========================================
FOLDER_ID = os.getenv("YANDEX_FOLDER_ID")
API_KEY = os.getenv("YANDEX_API_KEY")

if not FOLDER_ID or not API_KEY:
    raise ValueError("YANDEX_FOLDER_ID and YANDEX_API_KEY must be set in environment variables")
# ==========================================

YANDEXGPT_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"

# Коэффициент пересчёта слов в токены (для русского языка)
WORD_TO_TOKEN_RATIO = 6.0


# ----------------------------------------------------------------------
# Вспомогательные функции
# ----------------------------------------------------------------------
def count_words(text):
    """Точный подсчёт слов (буквенные последовательности, включая дефисы)."""
    words = re.findall(r'\b\w+(?:-\w+)*\b', text, flags=re.UNICODE)
    return len(words)


def make_summary_response(original_text, summary_text, max_tokens=None, desired_words=None, compression_percent=None):
    """Формирует JSON-ответ с метриками."""
    original_words = count_words(original_text)
    summary_words = count_words(summary_text)
    compression = round((1 - summary_words / original_words) * 100, 1) if original_words > 0 else 0
    return {
        'summary': summary_text,
        'original_words': original_words,
        'summary_words': summary_words,
        'compression': compression,
        'max_tokens_used': max_tokens,
        'desired_words': desired_words,
        'compression_percent': compression_percent
    }


# ----------------------------------------------------------------------
# Функция суммаризации
# ----------------------------------------------------------------------
def summarize_text(text, max_tokens=2000, desired_words=150, compression_percent=50, temperature=0.1, min_len=50):
    if not text:
        return "Текст пуст"

    if len(text) < min_len:
        logging.info("Текст слишком короткий, возвращаем как есть")
        return text

    max_input_chars = 15000
    if len(text) > max_input_chars:
        text = text[:max_input_chars] + "..."
        logging.info(f"Текст обрезан до {max_input_chars} символов")

    if compression_percent <= 30:
        style = "максимально краткий пересказ, выдели только самую суть"
    elif compression_percent <= 60:
        style = "сжатый пересказ, сохранив все ключевые моменты"
    else:
        style = "подробный пересказ с сохранением деталей"

    def make_request(extra_hint=""):
        length_hint = f"Сделай {style}. Твой ответ должен содержать ровно {desired_words} слов. Не останавливайся, пока не достигнешь этого числа. Добавляй детали, если нужно, чтобы заполнить объём. {extra_hint}"
        system_prompt = (
            "Ты – ассистент, который кратко и ёмко пересказывает тексты на русском языке. "
            "Сохраняй главную мысль, убирай воду. " + length_hint
        )
        payload = {
            "modelUri": f"gpt://{FOLDER_ID}/yandexgpt",
            "completionOptions": {
                "stream": False,
                "temperature": temperature,
                "maxTokens": max_tokens  # без запаса
            },
            "messages": [
                {"role": "system", "text": system_prompt},
                {"role": "user", "text": text}
            ]
        }
        headers = {
            "Authorization": f"Api-Key {API_KEY}",
            "x-folder-id": FOLDER_ID,
            "Content-Type": "application/json"
        }
        try:
            response = requests.post(YANDEXGPT_URL, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            result = response.json()
            summary = result['result']['alternatives'][0]['message']['text'].strip()
            return summary
        except Exception as e:
            logging.error(f"Ошибка в make_request: {e}")
            return f"Ошибка при генерации: {str(e)}"

    # Первый запрос
    summary = make_request("Старайся точно уложиться в это число.")
    if summary is None or not isinstance(summary, str):
        summary = "Не удалось получить ответ от модели."

    word_count = len(summary.split())
    #Логирование
    logging.info(f"Первый ответ: {word_count} слов, желаемое: {desired_words}")

    # Если ответ слишком короткий и это не ошибка, пробуем снова
    if (word_count < desired_words * 0.8 and desired_words > 20
            and not summary.startswith("Ошибка") and not summary.startswith("Не удалось")):
        logging.info(f"Первый ответ слишком короткий, пробуем повторный запрос...")
        summary2 = make_request("Это очень важно! Не давай слишком краткий ответ. Постарайся максимально приблизиться к указанному числу слов.")
        if summary2 is not None and isinstance(summary2, str):
            summary = summary2
            # Обновляем количество слов для дальнейшей обработки
            word_count = len(summary.split())
            logging.info(f"Второй ответ: {word_count} слов, желаемое: {desired_words}")

    # Обработка для малых значений (если желаемое < 20)
    if desired_words < 20:
        sentences = summary.split('. ')
        if len(sentences) > 1:
            summary = sentences[0] + '.'
        words = summary.split()
        if len(words) > desired_words * 1.5:
            summary = ' '.join(words[:desired_words]) + '...'

    return summary


# ----------------------------------------------------------------------
# Функции извлечения текста из разных источников
# ----------------------------------------------------------------------
def extract_text_from_url(url):
    """Извлекает основной текст из веб-страницы."""
    try:
        article = Article(url, language='ru')
        article.download()
        article.parse()
        text = article.text
        if text and len(text) > 100:
            logging.info(f"newspaper3k извлёк {len(text)} символов")
            return text
    except Exception as e:
        logging.warning(f"newspaper3k не сработал: {e}")

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
            element.decompose()
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='mw-parser-output')
        if main_content:
            text = main_content.get_text(separator='\n')
        else:
            text = soup.get_text(separator='\n')
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        if text and len(text) > 100:
            logging.info(f"BeautifulSoup извлёк {len(text)} символов")
            return text
    except Exception as e:
        logging.warning(f"requests+BeautifulSoup не сработал: {e}")

    raise ValueError("Не удалось извлечь текст из URL")


def extract_text_from_file(file):
    """Извлекает текст из загруженного файла (TXT, PDF, DOCX)."""
    filename = file.filename
    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'txt':
        return file.read().decode('utf-8')
    elif ext == 'pdf':
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
        return text
    elif ext == 'docx':
        doc = docx.Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text
    else:
        raise ValueError('Неподдерживаемый формат файла')


def extract_text_from_image(file, ocr_lang='ru', max_size=(1920, 1080)):
    """Извлекает текст из изображения с помощью EasyOCR."""
    reader = easyocr.Reader([ocr_lang, 'en'])
    image = Image.open(file)
    image.thumbnail(max_size, Image.Resampling.LANCZOS)
    if image.mode != 'RGB':
        image = image.convert('RGB')
    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp:
        image.save(tmp.name, 'JPEG', quality=85)
        tmp_path = tmp.name
    result = reader.readtext(tmp_path)
    os.unlink(tmp_path)
    text = ' '.join([item[1] for item in result])
    return text


# ----------------------------------------------------------------------
# Маршруты Flask
# ----------------------------------------------------------------------
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/summarize_text', methods=['POST'])
def summarize_text_route():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()

        # Получаем процент сжатия (от 1 до 100)
        compression_percent = data.get('compression_percent', 50)
        if not isinstance(compression_percent, (int, float)) or compression_percent < 1:
            compression_percent = 50
        if compression_percent > 100:
            compression_percent = 100

        if not text:
            return jsonify({'error': 'Введите текст'}), 400
        if len(text) < 20:
            return jsonify({'error': 'Текст слишком короткий (минимум 20 символов)'}), 400

        # Вычисляем желаемое число слов на основе процента
        original_word_count = count_words(text)
        desired_words = max(5, int(original_word_count * compression_percent / 100))
        max_tokens = int(desired_words * WORD_TO_TOKEN_RATIO)

        summary = summarize_text(text, max_tokens=max_tokens, desired_words=desired_words, compression_percent=compression_percent)
        if summary.startswith("Ошибка"):
            return jsonify({'error': summary}), 500

        return jsonify(make_summary_response(
            text, summary,
            max_tokens=max_tokens,
            desired_words=desired_words,
            compression_percent=compression_percent
        ))
    except Exception as e:
        logging.error(f"Ошибка в /summarize_text: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/summarize_url', methods=['POST'])
def summarize_url_route():
    try:
        data = request.get_json()
        url = data.get('url', '').strip()

        compression_percent = data.get('compression_percent', 50)
        if not isinstance(compression_percent, (int, float)) or compression_percent < 1:
            compression_percent = 50
        if compression_percent > 100:
            compression_percent = 100

        if not url:
            return jsonify({'error': 'Введите URL'}), 400

        text = extract_text_from_url(url)
        if not text or len(text) < 50:
            return jsonify({'error': 'Не удалось извлечь достаточный текст из URL'}), 400

        original_word_count = count_words(text)
        desired_words = max(5, int(original_word_count * compression_percent / 100))
        max_tokens = int(desired_words * WORD_TO_TOKEN_RATIO)

        summary = summarize_text(text, max_tokens=max_tokens, desired_words=desired_words, compression_percent=compression_percent)
        if summary.startswith("Ошибка"):
            return jsonify({'error': summary}), 500

        return jsonify(make_summary_response(
            text, summary,
            max_tokens=max_tokens,
            desired_words=desired_words,
            compression_percent=compression_percent
        ))
    except Exception as e:
        logging.error(f"Ошибка в /summarize_url: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/summarize_file', methods=['POST'])
def summarize_file_route():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Файл не загружен'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Файл не выбран'}), 400

        compression_percent = request.form.get('compression_percent', type=int, default=50)
        if compression_percent < 1:
            compression_percent = 50
        if compression_percent > 100:
            compression_percent = 100

        filename = file.filename
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

        image_extensions = {'jpg', 'jpeg', 'png', 'bmp', 'tiff'}
        if ext in image_extensions:
            text = extract_text_from_image(file)
        else:
            text = extract_text_from_file(file)

        if not text or len(text) < 50:
            return jsonify({'error': 'Из файла не удалось извлечь достаточный текст'}), 400

        original_word_count = count_words(text)
        desired_words = max(5, int(original_word_count * compression_percent / 100))
        max_tokens = int(desired_words * WORD_TO_TOKEN_RATIO)

        summary = summarize_text(text, max_tokens=max_tokens, desired_words=desired_words, compression_percent=compression_percent)
        if summary.startswith("Ошибка"):
            return jsonify({'error': summary}), 500

        return jsonify(make_summary_response(
            text, summary,
            max_tokens=max_tokens,
            desired_words=desired_words,
            compression_percent=compression_percent
        ))
    except Exception as e:
        logging.error(f"Ошибка в /summarize_file: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/download_txt', methods=['POST'])
def download_txt():
    try:
        data = request.get_json()
        summary = data.get('summary', '').strip()
        if not summary:
            return jsonify({'error': 'Нет текста для скачивания'}), 400
        txt_data = summary.encode('utf-8')
        return send_file(
            io.BytesIO(txt_data),
            mimetype='text/plain',
            as_attachment=True,
            download_name='summary.txt'
        )
    except Exception as e:
        logging.error(f"Ошибка в /download_txt: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/download_docx', methods=['POST'])
def download_docx():
    try:
        data = request.get_json()
        summary = data.get('summary', '').strip()
        if not summary:
            return jsonify({'error': 'Нет текста для скачивания'}), 400
        doc = DocxDocument()
        doc.add_heading('Краткое изложение', level=1)
        doc.add_paragraph(summary)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='summary.docx'
        )
    except Exception as e:
        logging.error(f"Ошибка в /download_docx: {e}")
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)